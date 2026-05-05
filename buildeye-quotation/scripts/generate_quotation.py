#!/usr/bin/env python3
"""BuildEye Quotation Generator - uses python-docx to fill the template."""
import argparse, copy, io, json, os, re, sys, shutil, tempfile
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import lxml.etree as etree

SERVICE_KEYWORDS = {
    "סריקה":   "סריקה (ענן נקודות)",
    "מודל":    "מודל תלת ממדי",
    "צפיין":   "צפיין בילדאיי",
    "autocad": "AutoCAD",
    "חזיתות":  "חזיתות",
    "חתכים":   "חתכים",
    "סיור":    "סיור וירטואלי",
}
LABEL_TO_KEY = {
    "סריקה": "סריקה", "מודל": "מודל", "צפיין": "צפיין",
    "תוכנית מדידה": "autocad", "autocad": "autocad",
    "חזיתות": "חזיתות", "חתכים": "חתכים",
    "סיור": "סיור", "סיור וירטואלי": "סיור",
    "סיור וירטואלי תלת-ממדי": "סיור",
}


def parse_services(s):
    keys = set()
    for label in (s or "").split(","):
        label = label.strip()
        key = LABEL_TO_KEY.get(label, label.lower())
        if key in SERVICE_KEYWORDS:
            keys.add(key)
        else:
            for k, v in LABEL_TO_KEY.items():
                if k in label or label in k:
                    if v in SERVICE_KEYWORDS:
                        keys.add(v); break
    return keys


def remove_row(row):
    tr = row._tr
    tr.getparent().remove(tr)


def copy_rpr(source_run_el):
    rPr = source_run_el.find(qn("w:rPr"))
    if rPr is not None:
        return copy.deepcopy(rPr)
    return None


def make_rtl_rpr():
    rPr = OxmlElement("w:rPr")
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return rPr


def make_heading1_title_rpr():
    """Times New Roman 18pt Bold Underline for quotation subject (הנדון)."""
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(rFonts)
    b = OxmlElement("w:b")
    rPr.append(b)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "36")  # 18pt = 36 half-points
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "36")
    rPr.append(szCs)
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return rPr


def make_arial11_rpr():
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "22")
    rPr.append(szCs)
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return rPr


def set_para_text_rtl(para, text):
    p = para._p
    existing_runs = p.findall(qn("w:r"))
    saved_rPr = copy_rpr(existing_runs[0]) if existing_runs else None
    for r in existing_runs:
        p.remove(r)
    run_el = OxmlElement("w:r")
    if saved_rPr is not None:
        run_el.append(saved_rPr)
    else:
        run_el.append(make_rtl_rpr())
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    run_el.append(t)
    p.append(run_el)


def set_recipient_two_lines(para, name, position_company):
    p = para._p
    pPr = p.find(qn("w:pPr"))
    base_rPr = None
    if pPr is not None:
        base_rPr = pPr.find(qn("w:rPr"))
    for r in p.findall(qn("w:r")):
        p.remove(r)

    def make_run(text, rpr_template=None):
        run_el = OxmlElement("w:r")
        if rpr_template is not None:
            run_el.append(copy.deepcopy(rpr_template))
        else:
            run_el.append(make_rtl_rpr())
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        run_el.append(t)
        return run_el

    def make_linebreak():
        run_el = OxmlElement("w:r")
        br = OxmlElement("w:br")
        run_el.append(br)
        return run_el

    p.append(make_run(name, base_rPr))
    p.append(make_linebreak())
    p.append(make_run(position_company, base_rPr))


def append_run_arial11(para, text):
    run_el = OxmlElement("w:r")
    run_el.append(make_arial11_rpr())
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    run_el.append(t)
    para._p.append(run_el)


def insert_para_after(ref_para, text):
    """Insert a new RTL paragraph after ref_para.
    No explicit bidi — sectPr+Normal style handle RTL direction.
    No explicit jc — RTL default is right-aligned."""""
    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = make_rtl_rpr()
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    new_p.append(r)
    ref_para._p.addnext(new_p)


def fill_sdt_arial11(para, value):
    sdt = para._p.find(qn("w:sdt"))
    if sdt is not None:
        sdtPr = sdt.find(qn("w:sdtPr"))
        if sdtPr is not None:
            ph = sdtPr.find(qn("w:showingPlcHdr"))
            if ph is not None:
                sdtPr.remove(ph)
        content_el = sdt.find(qn("w:sdtContent"))
        if content_el is not None:
            for r in content_el.iter(qn("w:r")):
                existing_rpr = r.find(qn("w:rPr"))
                if existing_rpr is not None:
                    r.remove(existing_rpr)
                r.insert(0, make_arial11_rpr())
                for t_el in r.iter(qn("w:t")):
                    t_el.text = value
                    break
                break
    else:
        run = para.add_run(value)
        run.font.name = "Arial"
        run.font.size = Pt(11)


def _insert_before_rpr(pPr, element):
    """Insert element into pPr before w:rPr (which must always be last in pPr).
    If no rPr exists, appends at end."""
    rPr = pPr.find(qn("w:rPr"))
    if rPr is not None:
        idx = list(pPr).index(rPr)
        pPr.insert(idx, element)
    else:
        pPr.append(element)


def ensure_rtl_all_paragraphs(doc):
    """Post-process: add w:bidi to paragraphs that are missing it.
    Preserves existing jc values (center, right, etc.) — only adds jc=right
    when no jc is present AND no paragraph style is set (bare paragraphs added
    by the script that wouldn't inherit the Normal-style bidi).
    The Normal style already has w:bidi, so most paragraphs are fine without
    a paragraph-level bidi — we only add it where truly missing."""
    def fix_para(para):
        p = para._p
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        bidi_el = pPr.find(qn("w:bidi"))
        jc_el   = pPr.find(qn("w:jc"))
        pStyle_el = pPr.find(qn("w:pStyle"))

        # Add bidi before rPr if missing
        if bidi_el is None:
            bidi_el = OxmlElement("w:bidi")
            _insert_before_rpr(pPr, bidi_el)

        # Add jc=right whenever bidi is set and no explicit jc exists.
        # Never override center — table header cells are intentionally centered.
        # Flip explicit "left" to "right" (shouldn't happen in Hebrew doc but be safe).
        if jc_el is None:
            jc_el = OxmlElement("w:jc")
            jc_el.set(qn("w:val"), "right")
            _insert_before_rpr(pPr, jc_el)
        elif jc_el.get(qn("w:val")) == "left":
            jc_el.set(qn("w:val"), "right")
        # jc=center → keep as-is (intentional centering in table headers etc.)

    for para in doc.paragraphs:
        fix_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fix_para(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            fix_para(para)
        for para in section.footer.paragraphs:
            fix_para(para)
    print("  [OK] RTL pass: bidi added where missing, center alignment preserved")


def renumber_hash_column(table):
    if len(table.rows) < 2:
        return
    hash_col_idx = None
    header_cells = table.rows[0].cells
    for i, cell in enumerate(header_cells):
        if cell.text.strip() == '#':
            hash_col_idx = i
            break
    if hash_col_idx is None:
        for col_idx in range(len(header_cells)):
            try:
                val = table.rows[1].cells[col_idx].text.strip()
                if val.isdigit():
                    hash_col_idx = col_idx
                    break
            except IndexError:
                pass
    if hash_col_idx is None:
        print("  [WARN] Could not find '#' column for renumbering")
        return
    for num, row in enumerate(table.rows[1:], 1):
        if hash_col_idx < len(row.cells):
            cell = row.cells[hash_col_idx]
            for para in cell.paragraphs:
                set_para_text_rtl(para, str(num))
    print(f"  [OK] Renumbered '#' column (col {hash_col_idx}), {len(table.rows)-1} rows")


def replace_text_in_doc(doc, old, new):
    count = 0
    for para in doc.paragraphs:
        for t_el in para._p.iter(qn("w:t")):
            if t_el.text and old in t_el.text:
                t_el.text = t_el.text.replace(old, new)
                count += 1
    return count


def fill_header_quote_number(doc, quote_number):
    for section in doc.sections:
        hdr = section.header
        for para in hdr.paragraphs:
            if para._p.get(qn("w14:paraId")) == "4D8FB4BF":
                p = para._p
                runs = p.findall(qn("w:r"))
                for r in runs:
                    t = r.find(qn("w:t"))
                    if t is not None and (t.text is None or t.text.strip() == ""):
                        t.text = " " + str(quote_number)
                        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        print(f"  [OK] Quote number in header: {quote_number}")
                        return
                run_el = OxmlElement("w:r")
                run_el.append(make_rtl_rpr())
                t = OxmlElement("w:t")
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t.text = " " + str(quote_number)
                run_el.append(t)
                p.append(run_el)
                print(f"  [OK] Quote number in header (appended): {quote_number}")
                return


def get_and_increment_quote_number(output_dir):
    counter_file = os.path.join(output_dir, "quote_counter.json")
    if os.path.exists(counter_file):
        with open(counter_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        current = data.get('next_quote_number', 321)
    else:
        current = 321
    with open(counter_file, 'w', encoding='utf-8') as f:
        json.dump({'next_quote_number': current + 1}, f, ensure_ascii=False)
    return current


def generate_output_name(title, output_dir, quote_number, company_name):
    name = title
    for prefix in ["הצעת מחיר לסריקה ומידול ", "הצעת מחיר ל", "הצעת מחיר - ", "הצעת מחיר"]:
        if name.startswith(prefix):
            name = name[len(prefix):]
            break
    safe_company = re.sub(r'[<>:"/\\|?*]', "-", company_name).strip()
    safe_project = re.sub(r'[<>:"/\\|?*]', "-", name).strip()
    filename = f"{quote_number}_{safe_company} - הצעת מחיר לסריקה ומידול {safe_project}.docx"
    return os.path.join(output_dir, filename)


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--template", required=True)
    p.add_argument("--output-dir", required=True)
    p.add_argument("--counter-dir", default=None,
                   help="Directory holding quote_counter.json (defaults to output-dir)")
    p.add_argument("--client-name", required=True)
    p.add_argument("--client-position", required=True)
    p.add_argument("--company-name", required=True)
    p.add_argument("--title", required=True)
    p.add_argument("--project", required=True)
    p.add_argument("--services", required=True)
    p.add_argument("--pricing-notes", default="")
    p.add_argument("--scale", default="")
    p.add_argument("--delivery-days", required=True)
    p.add_argument("--area-basis", required=True)
    p.add_argument("--payment-terms", required=True)
    a = p.parse_args()

    if not os.path.isfile(a.template):
        print("ERROR: Template not found: " + a.template, file=sys.stderr); sys.exit(1)

    selected = parse_services(a.services)
    has_autocad = "autocad" in selected
    has_model   = "מודל" in selected
    mapping_mode = has_autocad and not has_model
    print("  Services: " + str(selected))
    print("  Mapping mode (AutoCAD, no 3D model): " + str(mapping_mode))

    counter_dir = a.counter_dir if a.counter_dir else a.output_dir
    quote_number = get_and_increment_quote_number(counter_dir)
    print(f"  Quote number: {quote_number}")

    doc = Document(a.template)
    paragraphs = doc.paragraphs

    # 1. Recipient
    for para in paragraphs:
        if para._p.get(qn("w14:paraId")) == "4BF2AFF9":
            position_company = f"{a.client_position}, {a.company_name}"
            set_recipient_two_lines(para, a.client_name, position_company)
            print(f"  [OK] Recipient: {a.client_name} | {position_company}")
            break

    # 2. Heading — Heading 1 style, Times New Roman 18 Bold Underline
    for para in paragraphs:
        if "לסריקה ומידול" in para.text and "הנדון" in para.text:
            p_el = para._p
            # Apply Heading 1 paragraph style
            pPr = p_el.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p_el.insert(0, pPr)
            pStyle_el = pPr.find(qn("w:pStyle"))
            if pStyle_el is None:
                pStyle_el = OxmlElement("w:pStyle")
                pPr.insert(0, pStyle_el)
            pStyle_el.set(qn("w:val"), "1")
            # Replace runs with Times New Roman 18 Bold Underline
            for r in p_el.findall(qn("w:r")):
                p_el.remove(r)
            run_el = OxmlElement("w:r")
            run_el.append(make_heading1_title_rpr())
            t = OxmlElement("w:t")
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = "הנדון: " + a.title
            run_el.append(t)
            p_el.append(run_el)
            print("  [OK] Heading (Heading1, TNR 18 Bold Underline): " + a.title)
            break

    # 3. Body text
    for para in paragraphs:
        if "_________" in para.text:
            new_text = para.text.replace("_________", a.project)
            set_para_text_rtl(para, new_text)
            print("  [OK] Body: " + a.project)
            break

    # 4. Pricing notes
    if a.pricing_notes and a.pricing_notes.strip():
        for para in paragraphs:
            if 'סה"כ לסריקה ומידול' in para.text:
                insert_para_after(para, a.pricing_notes)
                print('  [OK] Pricing notes (after סה"כ): ' + a.pricing_notes)
                break
    else:
        print("  [OK] Pricing notes - skipped")

    # 5. Scale
    for para in paragraphs:
        if 'קנ"מ של' in para.text or 'קנ\\"מ של' in para.text:
            if has_autocad and a.scale:
                append_run_arial11(para, a.scale)
                print("  [OK] Scale (Arial 11): " + a.scale)
            else:
                para._p.getparent().remove(para._p)
                print("  [OK] Removed scale paragraph")
            break

    # 6. Delivery days
    for para in paragraphs:
        if "______" in para.text and "ימי עסקים" in para.text:
            new_text = para.text.replace("______", a.delivery_days)
            set_para_text_rtl(para, new_text)
            print("  [OK] Delivery days: " + a.delivery_days)
            break

    # 7. Area basis
    for para in paragraphs:
        if "יכלול את" in para.text:
            fill_sdt_arial11(para, a.area_basis)
            print("  [OK] Area basis (Arial 11): " + a.area_basis)
            break

    # 8. Payment terms
    for para in paragraphs:
        if "תנאי התשלום" in para.text:
            fill_sdt_arial11(para, a.payment_terms)
            print("  [OK] Payment terms (Arial 11): " + a.payment_terms)
            break

    # 9. AutoCAD-without-model adjustments
    if has_autocad and not has_model:
        for para in list(paragraphs):
            if "Revit" in para.text and "2023" in para.text:
                para._p.getparent().remove(para._p)
                print("  [OK] Removed Revit paragraph")
                break
        for para in list(paragraphs):
            if "ההצעה אינה כוללת המרה של המודל התלת ממדי" in para.text:
                para._p.getparent().remove(para._p)
                print("  [OK] Removed 'המרה' line")
                break
        for old, new in [("ימודל", "ימופה"), ("מידול", "מיפוי")]:
            n = replace_text_in_doc(doc, old, new)
            if n:
                print(f"  [OK] Replaced '{old}'→'{new}' ({n} occurrences)")

    # 10. Remove unselected service rows
    table = doc.tables[0]
    rows_to_remove = []
    for row in table.rows[1:]:
        if len(row.cells) < 2:
            continue
        service_cell = row.cells[1].text.strip()
        for key, keyword in SERVICE_KEYWORDS.items():
            if keyword in service_cell and key not in selected:
                rows_to_remove.append(row)
                print("  [OK] Removing row: " + service_cell[:30])
                break
    for row in rows_to_remove:
        remove_row(row)

    # 11. Renumber '#' column
    renumber_hash_column(table)

    # 12. Fill header quote number
    fill_header_quote_number(doc, quote_number)

    # 13. RTL pass — disabled: sectPr bidi + Normal style bidi handle it
    # ensure_rtl_all_paragraphs(doc)

    os.makedirs(a.output_dir, exist_ok=True)
    # Save via BytesIO then write directly — avoids temp-file copy corruption
    output_path = generate_output_name(a.title, a.output_dir, quote_number, a.company_name)
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    with open(output_path, "wb") as f:
        f.write(content)
        f.flush()
        os.fsync(f.fileno())
    print("SUCCESS: " + output_path)
    print("QUOTE_NUMBER: " + str(quote_number))


if __name__ == "__main__":
    main()
